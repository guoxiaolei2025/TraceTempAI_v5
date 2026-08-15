import logging
import re
from datetime import datetime, timedelta
from typing import Dict, Optional

from core.data_collector import DataCollector
from core.report_generator import ReportGenerator
from core.task_manager import TaskManager
from core.department_manager import DepartmentManager
from core.ai_analyzer import AIAnalyzer

logger = logging.getLogger(__name__)

# AI 报告免责声明（追加于 AI 生成的分析文本末尾）
AI_REPORT_DISCLAIMER = "该报告由AI自动生成，可能存在分析或描述错误，仅供参考使用。"

# 晦涩英文缩写 -> 中文全称（过滤表，可扩展）
ABBREVIATION_MAP = {
    "FMEA": "设备故障模式与影响分析",
    "HVAC": "空调通风系统",
    "SOP": "标准操作规程",
    "UPS": "不间断电源",
}

class AlarmService:
    def __init__(self):
        self.dept_manager = DepartmentManager()
        self.ai_analyzer = AIAnalyzer()
    
    def get_alarm_data(self, start_date: str, end_date: str, dept_id: Optional[str] = None):
        if dept_id:
            dept = self.dept_manager.get_department(dept_id)
            if not dept:
                raise ValueError(f"学科 {dept_id} 不存在")
            api_key = dept.api_key
        else:
            departments = self.dept_manager.get_all_departments()
            if not departments:
                raise ValueError("未配置任何学科")
            api_key = departments[0].api_key
        
        if not api_key:
            raise ValueError("学科API Key未配置")
        
        start_dt = datetime.strptime(start_date, "%Y-%m-%d")
        end_dt = datetime.strptime(end_date, "%Y-%m-%d").replace(hour=23, minute=59, second=59)
        
        collector = DataCollector(api_key=api_key)
        all_alarms = collector.get_all_alarms(start_dt, end_dt)
        
        total = 0
        handled_count = 0
        unhandled_count = 0
        devices_list = []
        
        for device_name, alarms in all_alarms.items():
            for alarm in alarms:
                total += 1
                if alarm["handlestate"] == "已处理":
                    handled_count += 1
                else:
                    unhandled_count += 1
            
            device_parts = device_name.split("【")
            device_name_str = device_parts[0]
            device_sn = device_parts[1].replace("】", "") if len(device_parts) > 1 else ""
            devices_list.append({"device_name": device_name_str, "device_sn": device_sn, "alarms": alarms})
        
        return {
            "total": total,
            "handled_count": handled_count,
            "unhandled_count": unhandled_count,
            "devices": devices_list,
            "statistics": {
                "total_alarms": total,
                "temperature_alarms": total // 2,
                "humidity_alarms": total // 2,
                "device_count": len(devices_list)
            }
        }
    
    def analyze_alarms(self, start_date: str, end_date: str, alarm_data: Dict, dept_id: Optional[str] = None):
        total_alarms = alarm_data.get("total_alarms", 0)
        handled_count = alarm_data.get("handled_count", 0)
        unhandled_count = alarm_data.get("unhandled_count", 0)
        handled_rate = round((handled_count / total_alarms * 100), 2) if total_alarms > 0 else 100
        
        devices = alarm_data.get("devices", [])
        
        # 1. 设备报警排行
        device_alarm_counts = []
        for device in devices:
            device_name = device.get("device_name", "")
            device_sn = device.get("device_sn", "")
            alarm_count = len(device.get("alarms", []))
            full_name = f"{device_name}【{device_sn}】" if device_sn else device_name
            device_alarm_counts.append({
                "name": full_name,
                "alarm_count": alarm_count,
                "device": device
            })
        
        device_alarm_counts.sort(key=lambda x: x["alarm_count"], reverse=True)
        top_devices = []
        for i, d in enumerate(device_alarm_counts[:5]):
            top_devices.append({
                "name": d["name"],
                "alarm_count": d["alarm_count"],
                "rank": i + 1
            })
        
        # 2. 按类型分析：区分冰箱温度和环境温湿度
        fridge_devices = []
        env_devices = []
        
        for d in device_alarm_counts:
            device_name = d["name"]
            if "冰箱" in device_name or "冷藏" in device_name or "冷冻" in device_name:
                fridge_devices.append(d)
            else:
                env_devices.append(d)
        
        # 冰箱温度报警排行（重点）
        fridge_top_devices = []
        for i, d in enumerate(fridge_devices[:5]):
            fridge_top_devices.append({
                "name": d["name"],
                "alarm_count": d["alarm_count"],
                "rank": i + 1
            })
        
        # 环境温湿度报警排行
        env_top_devices = []
        for i, d in enumerate(env_devices[:5]):
            env_top_devices.append({
                "name": d["name"],
                "alarm_count": d["alarm_count"],
                "rank": i + 1
            })
        
        # 3. 报警类型统计 - 统计温度、湿度、其他
        temp_count = 0
        hum_count = 0
        other_count = 0
        
        hour_data = {h: 0 for h in range(24)}
        week_data = {i: 0 for i in range(7)}
        month_data = {i: 0 for i in range(1, 32)}  # 月度1-31日
        
        # 收集未处理报警详情和响应时间
        unhandled_details = []
        response_times = []  # 响应时间列表（分钟）
        
        for device in devices:
            device_name = device.get("device_name", "")
            device_sn = device.get("device_sn", "")
            full_name = f"{device_name}【{device_sn}】" if device_sn else device_name
            for alarm in device.get("alarms", []):
                msg = alarm.get("message", "")
                has_temp = "温度" in msg
                has_hum = "湿度" in msg
                if has_temp:
                    temp_count += 1
                if has_hum:
                    hum_count += 1
                if not has_temp and not has_hum:
                    other_count += 1
                
                # 统计时段、星期、日期
                alarm_dt = None
                try:
                    alarm_dt = datetime.fromisoformat(alarm.get("alarmdate", "").replace("Z", "+00:00"))
                    hour_data[alarm_dt.hour] += 1
                    week_data[alarm_dt.weekday()] += 1
                    month_data[alarm_dt.day] += 1
                except:
                    pass
                
                # 收集未处理报警
                if alarm.get("handlestate") != "已处理":
                    unhandled_details.append({
                        "device": full_name,
                        "alarm_time": alarm.get("alarmdate", ""),
                        "time": alarm.get("alarmdate", ""),
                        "message": msg,
                        "type": "温度" if has_temp and not has_hum else ("湿度" if has_hum and not has_temp else "温湿度"),
                        "severity": "高"
                    })
                
                # 计算响应时间（如果有处理时间信息）
                if alarm.get("handlestate") == "已处理" and alarm_dt:
                    handle_time_str = alarm.get("handletime", "")
                    if handle_time_str:
                        try:
                            handle_dt = datetime.fromisoformat(handle_time_str.replace("Z", "+00:00"))
                            diff_minutes = (handle_dt - alarm_dt).total_seconds() / 60
                            if 0 < diff_minutes < 10080:  # 排除异常值（>7天）
                                response_times.append(diff_minutes)
                        except:
                            pass
        
        # 计算平均响应时间
        if response_times:
            avg_minutes = sum(response_times) / len(response_times)
            hours = int(avg_minutes // 60)
            minutes = int(avg_minutes % 60)
            average_response_time = f"{hours}小时{minutes}分钟" if hours > 0 else f"{minutes}分钟"
        else:
            average_response_time = "暂无数据"
        
        # 分析高峰时段 - 找连续的高峰时段范围
        peak_hours = []
        max_hour_count = max(hour_data.values()) if hour_data.values() else 0
        if max_hour_count > 0:
            # 找出所有高峰小时
            peak_hour_list = []
            for h, count in hour_data.items():
                if count >= max_hour_count * 0.8:  # 包含前80%的高峰
                    peak_hour_list.append(h)
            
            if peak_hour_list:
                # 找连续的时段
                peak_hour_list.sort()
                start_h = peak_hour_list[0]
                end_h = peak_hour_list[-1]
                
                if start_h == end_h:
                    peak_hours.append(f"{start_h}:00-{start_h+1}:00")
                else:
                    peak_hours.append(f"{start_h}:00-{end_h+1}:00")
        
        
        # 分析高峰星期
        weekdays = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"]
        peak_days = []
        max_day_count = max(week_data.values()) if week_data.values() else 0
        if max_day_count > 0:
            for i, count in week_data.items():
                if count == max_day_count:
                    peak_days.append(weekdays[i])
                    if len(peak_days) >= 2:
                        break
        
        # 分析月度日期高峰
        peak_month_days = []
        max_month_day_count = max(month_data.values()) if month_data.values() else 0
        if max_month_day_count > 0:
            # 找连续的高峰日期
            peak_day_list = []
            for d, count in month_data.items():
                if count >= max_month_day_count * 0.8:
                    peak_day_list.append(d)
            
            if peak_day_list:
                peak_day_list.sort()
                start_d = peak_day_list[0]
                end_d = peak_day_list[-1]
                
                if start_d == end_d:
                    peak_month_days.append(f"{start_d}日")
                else:
                    peak_month_days.append(f"{start_d}日-{end_d}日")
        
        # 统计数据汇总
        type_total = temp_count + hum_count + other_count
        type_total = type_total if type_total > 0 else 1
        temp_percent = round(temp_count / type_total * 100, 2) if temp_count > 0 else 0
        hum_percent = round(hum_count / type_total * 100, 2) if hum_count > 0 else 0
        other_percent = round(other_count / type_total * 100, 2) if other_count > 0 else 0
        fridge_alarm_total = sum(d["alarm_count"] for d in fridge_devices)
        env_alarm_total = sum(d["alarm_count"] for d in env_devices)
        total_devices = len(devices)
        fridge_ratio = round(len(fridge_devices) / total_devices * 100, 2) if total_devices > 0 else 0
        fridge_alarm_ratio = round(fridge_alarm_total / total_alarms * 100, 2) if total_alarms > 0 else 0
        
        # 严重程度统计：按报警次数排序，前10%为高，10%-60%为中，60%以后为低
        severity_high = 0
        severity_medium = 0
        severity_low = 0
        if device_alarm_counts:
            total_device_count = len(device_alarm_counts)
            for idx, d in enumerate(device_alarm_counts):
                ratio = idx / total_device_count
                if ratio < 0.1:
                    severity_high += d["alarm_count"]
                elif ratio < 0.6:
                    severity_medium += d["alarm_count"]
                else:
                    severity_low += d["alarm_count"]
        
        # 4. 调用AI大模型进行深度分析
        ai_statistics = {
            'total_alarms': total_alarms,
            'handled_count': handled_count,
            'unhandled_count': unhandled_count,
            'start_date': start_date,
            'end_date': end_date,
            'top_devices': [{'name': d['name'], 'alarm_count': d['alarm_count']} for d in device_alarm_counts[:10]],
            'fridge_info': {
                'count': len(fridge_devices),
                'alarm_total': fridge_alarm_total
            },
            'env_info': {
                'count': len(env_devices),
                'alarm_total': env_alarm_total
            },
            'type_distribution': {
                '温度': temp_count,
                '湿度': hum_count,
                '其他': other_count
            },
            'peak_hours': peak_hours,
            'peak_days': peak_days,
            'peak_month_days': peak_month_days,
            'unhandled_details': unhandled_details[:10]
        }
        
        ai_result = self.ai_analyzer.analyze(ai_statistics)
        
        # 5. 生成建议（优先使用AI建议，降级使用规则建议）
        suggestions = []
        critical_devices = []
        
        if ai_result and ai_result.get('ai_suggestions'):
            # 使用AI生成的建议（规范化：建议开头、过滤英文缩写）
            suggestions = self._normalize_suggestions(ai_result['ai_suggestions'])
        else:
            # 降级：使用规则生成建议
            if top_devices:
                top_device = top_devices[0]
                critical_devices.append(top_device["name"])
                suggestions.append({
                    "type": "设备维护",
                    "target": top_device["name"],
                    "description": f"该设备报警次数最多（{top_device['alarm_count']}次），建议安排专业维护人员进行全面检查",
                    "priority": "高"
                })
            
            if fridge_top_devices:
                fridge_top = fridge_top_devices[0]
                if fridge_top["name"] != (top_devices[0]["name"] if top_devices else ""):
                    critical_devices.append(fridge_top["name"])
                    suggestions.append({
                        "type": "冰箱重点维护",
                        "target": fridge_top["name"],
                        "description": f"该冰箱设备报警次数最多（{fridge_top['alarm_count']}次），温度失控对试剂或样本影响重大，建议立即安排紧急检查和维护",
                        "priority": "高"
                    })
            
            if peak_hours:
                suggestions.append({
                    "type": "环境调整",
                    "target": f"{peak_hours[0]}前后",
                    "description": "该时段报警频率最高，建议检查空调运行状态和设备散热情况",
                    "priority": "中"
                })
            
            if peak_days:
                suggestions.append({
                    "type": "管理优化",
                    "target": peak_days[0],
                    "description": f"{peak_days[0]}报警次数明显高于其他日期，建议加强巡检频次",
                    "priority": "中"
                })
            
            suggestions.append({
                "type": "预防性维护",
                "target": "全部设备",
                "description": "建议制定季度预防性维护计划，尤其是冰箱设备应每月检查一次",
                "priority": "低"
            })
            suggestions = self._normalize_suggestions(suggestions)
        
        # 6. 生成分析文本（优先使用AI文本，降级使用本地模板）
        if ai_result and ai_result.get('analysis_text'):
            analysis_text = ai_result['analysis_text']
            # 过滤英文缩写并追加免责声明
            analysis_text = self._filter_abbreviations(analysis_text)
            if AI_REPORT_DISCLAIMER not in analysis_text:
                analysis_text = f"{analysis_text.rstrip()}\n\n说明：{AI_REPORT_DISCLAIMER}"
            trend_assessment = ai_result.get('trend_assessment', '')
        else:
            # 降级：使用本地规则生成分析文本
            analysis_text = self._generate_fallback_analysis(
                start_date, end_date, total_alarms, handled_rate,
                total_devices, temp_count, hum_count, other_count,
                temp_percent, hum_percent, other_percent,
                fridge_devices, env_devices, fridge_ratio, fridge_alarm_ratio,
                fridge_alarm_total, env_alarm_total,
                fridge_top_devices, env_top_devices,
                peak_hours, peak_days, peak_month_days
            )
            trend_assessment = ''

        # 6.1 计算环比变化（百分比）
        mom = self._compute_mom_change(start_date, end_date, total_alarms, dept_id)
        if mom:
            trend_change = mom['text']
            previous_total_alarms = mom['previous']
        else:
            trend_change = '暂无对比数据'
            previous_total_alarms = 0

        # 构建未处理报警列表（按设备聚合，避免同一设备占满前10条）
        unhandled_by_device = {}
        for item in unhandled_details:
            dev = item['device']
            atime = item.get('alarm_time') or ''
            if dev not in unhandled_by_device:
                unhandled_by_device[dev] = {
                    'device': dev,
                    'alarm_count': 1,
                    'first_alarm_time': atime,
                    'latest_alarm_time': atime,
                    'type': item.get('type', ''),
                    'severity': '高'
                }
            else:
                entry = unhandled_by_device[dev]
                entry['alarm_count'] += 1
                if atime:
                    if not entry['first_alarm_time'] or atime < entry['first_alarm_time']:
                        entry['first_alarm_time'] = atime
                    if atime > entry['latest_alarm_time']:
                        entry['latest_alarm_time'] = atime
        unhandled_list = sorted(
            unhandled_by_device.values(), key=lambda x: x['alarm_count'], reverse=True
        )[:10]
        for item in unhandled_list:
            # 兼容旧字段：alarm_time 用最早一条
            item['alarm_time'] = item['first_alarm_time']
        
        return {
            "overview": {
                "total_alarms": total_alarms,
                "handled_rate": handled_rate,
                "trend_change": trend_change,
                "trend_assessment": trend_assessment,
                "previous_total_alarms": previous_total_alarms,
                "period": f"{start_date} ~ {end_date}"
            },
            "device_analysis": {
                "top_devices": top_devices,
                "critical_devices": critical_devices if critical_devices else [d["name"] for d in top_devices[:2]],
                "peak_hours": peak_hours[:2],
                "peak_days": peak_days[:2],
                "peak_month_days": peak_month_days[:2]
            },
            "type_analysis": {
                "by_category": {
                    "fridge": {
                        "name": "冰箱温度",
                        "devices": fridge_devices,
                        "top_devices": fridge_top_devices,
                        "total_alarms": fridge_alarm_total
                    },
                    "environment": {
                        "name": "环境温湿度",
                        "devices": env_devices,
                        "top_devices": env_top_devices,
                        "total_alarms": env_alarm_total
                    }
                },
                "distribution": [
                    {"type": "温度", "count": temp_count, "percentage": temp_percent},
                    {"type": "湿度", "count": hum_count, "percentage": hum_percent},
                    {"type": "其他", "count": other_count, "percentage": other_percent}
                ],
                "severity_statistics": {"高": severity_high, "中": severity_medium, "低": severity_low}
            },
            "suggestions": suggestions,
            "handling_evaluation": {
                "unhandled_count": unhandled_count,
                "average_response_time": average_response_time,
                "unhandled_list": unhandled_list,
                "recommended_priority": [d["name"] for d in top_devices[:3]]
            },
            "analysis_text": analysis_text,
            "ai_powered": ai_result is not None,
            "quality_score": ai_result.get("quality_score", {}).get("total_score", 0) if ai_result else 0
        }
    
    @staticmethod
    def _filter_abbreviations(text: str) -> str:
        """将晦涩的英文缩写替换为中文全称"""
        if not text:
            return text
        for abbr, full in ABBREVIATION_MAP.items():
            # 中文属于 \w，不能使用 \b，改用字母数字边界匹配
            text = re.sub(rf'(?<![A-Za-z0-9]){abbr}(?![A-Za-z0-9])', full, text)
        # 修复 FMEA 映射为"…影响分析"后与后续"分析"重叠的情况
        text = text.replace('影响分析分析', '影响分析')
        return text

    @staticmethod
    def _normalize_suggestions(suggestions: list) -> list:
        """
        规范化建议列表：
        1. 每条 description 以"建议"开头（无则添加）
        2. 过滤 description 中的英文缩写
        """
        normalized = []
        for item in suggestions:
            if not isinstance(item, dict):
                continue
            desc = AlarmService._filter_abbreviations(item.get('description', '') or '')
            desc = desc.strip()
            if desc and not desc.startswith('建议'):
                # 去除句中冗余的"建议"字样，统一前置，避免"建议：…建议…"重复
                desc = re.sub(r'[，,；;]\s*建议', '，', desc)
                desc = re.sub(r'^\s*建议[：:，,]?\s*', '', desc).strip()
                desc = f'建议：{desc}'
            item['description'] = desc
            normalized.append(item)
        return normalized

    def _compute_mom_change(self, start_date: str, end_date: str, current_total: int, dept_id: Optional[str]) -> Optional[Dict]:
        """查询前一同长周期报警总量，计算环比变化百分比"""
        if not dept_id:
            return None
        try:
            start_dt = datetime.strptime(start_date, "%Y-%m-%d")
            end_dt = datetime.strptime(end_date, "%Y-%m-%d")
            period_days = (end_dt - start_dt).days + 1
            if period_days <= 0:
                return None
            prev_end_dt = start_dt - timedelta(days=1)
            prev_start_dt = prev_end_dt - timedelta(days=period_days - 1)
            prev_start = prev_start_dt.strftime("%Y-%m-%d")
            prev_end = prev_end_dt.strftime("%Y-%m-%d")
            prev_data = self.get_alarm_data(prev_start, prev_end, dept_id)
            prev_total = prev_data.get("total", 0)
            if prev_total == 0:
                if current_total == 0:
                    return {"text": "0%", "previous": 0}
                return {"text": "新增", "previous": 0}
            change_pct = (current_total - prev_total) / prev_total * 100
            sign = "+" if change_pct > 0 else ""
            return {"text": f"{sign}{change_pct:.1f}%", "previous": prev_total}
        except Exception as e:
            logger.warning(f"环比计算失败: {e}")
            return None

    def _generate_fallback_analysis(self, start_date, end_date, total_alarms, handled_rate,
                                     total_devices, temp_count, hum_count, other_count,
                                     temp_percent, hum_percent, other_percent,
                                     fridge_devices, env_devices, fridge_ratio, fridge_alarm_ratio,
                                     fridge_alarm_total, env_alarm_total,
                                     fridge_top_devices, env_top_devices,
                                     peak_hours, peak_days, peak_month_days):
        """降级生成本地分析文本（当AI服务不可用时使用）"""
        analysis_lines = []
        
        analysis_lines.append("【报警概况总结】")
        analysis_lines.append(f"时间范围：{start_date} 至 {end_date}")
        analysis_lines.append(f"报警总数：{total_alarms} 次")
        analysis_lines.append(f"处理完成率：{handled_rate}%")
        analysis_lines.append(f"监控设备数：{total_devices} 台")
        analysis_lines.append("")
        
        analysis_lines.append("【报警类型分析】")
        if temp_count > 0:
            analysis_lines.append(f"- 温度报警：{temp_count} 次，占比 {temp_percent}%")
        if hum_count > 0:
            analysis_lines.append(f"- 湿度报警：{hum_count} 次，占比 {hum_percent}%")
        if other_count > 0:
            analysis_lines.append(f"- 其他报警：{other_count} 次，占比 {other_percent}%")
        analysis_lines.append("")
        
        analysis_lines.append("【设备分类分析】")
        if fridge_devices:
            analysis_lines.append(f"■ 冰箱/冷藏设备：")
            analysis_lines.append(f"  - 设备数：{len(fridge_devices)} 台，占比 {fridge_ratio}%")
            analysis_lines.append(f"  - 报警数：{fridge_alarm_total} 次，占总报警 {fridge_alarm_ratio}%")
            analysis_lines.append(f"  - 平均每台：{round(fridge_alarm_total / len(fridge_devices), 1)} 次/台")
            if fridge_top_devices:
                analysis_lines.append(f"  - TOP设备：{fridge_top_devices[0]['name']}，报警 {fridge_top_devices[0]['alarm_count']} 次")
        
        if env_devices:
            analysis_lines.append(f"■ 环境温湿度设备：")
            analysis_lines.append(f"  - 设备数：{len(env_devices)} 台，占比 {round(100 - fridge_ratio, 2)}%")
            analysis_lines.append(f"  - 报警数：{env_alarm_total} 次，占总报警 {round(100 - fridge_alarm_ratio, 2)}%")
            analysis_lines.append(f"  - 平均每台：{round(env_alarm_total / len(env_devices), 1)} 次/台")
            if env_top_devices:
                analysis_lines.append(f"  - TOP设备：{env_top_devices[0]['name']}，报警 {env_top_devices[0]['alarm_count']} 次")
        analysis_lines.append("")
        
        analysis_lines.append("【时间规律分析】")
        if peak_hours:
            analysis_lines.append(f"- 报警高峰时段：{'、'.join(peak_hours[:2])}")
        else:
            analysis_lines.append("- 报警时段分布：全天相对均匀")
        if peak_days:
            analysis_lines.append(f"- 报警高峰星期：{'、'.join(peak_days[:2])}")
        else:
            analysis_lines.append("- 报警星期分布：全周相对均匀")
        if peak_month_days:
            analysis_lines.append(f"- 报警高峰日期：{'、'.join(peak_month_days[:2])}")
        else:
            analysis_lines.append("- 报警日期分布：全月相对均匀")
        analysis_lines.append("")
        
        analysis_lines.append("【客观评价与建议】")
        if fridge_alarm_ratio > 50:
            analysis_lines.append("⚠ 重点提示：冰箱设备报警占比超过50%，建议优先关注")
            analysis_lines.append("  原因：温度失控对试剂、样本等影响重大，风险较高")
        elif fridge_alarm_ratio > 30:
            analysis_lines.append("→ 提示：冰箱设备报警占比30%以上，建议关注")
        else:
            analysis_lines.append("→ 冰箱设备报警占比处于正常范围")
        
        if handled_rate < 70:
            analysis_lines.append("⚠ 重点提示：处理完成率低于70%，请检查报警响应机制")
        elif handled_rate < 90:
            analysis_lines.append("→ 提示：处理完成率尚有提升空间")
        else:
            analysis_lines.append("✓ 处理完成率良好")
        
        if peak_hours:
            analysis_lines.append(f"→ 建议在高峰时段 {'、'.join(peak_hours[:1])} 加强巡检和监控")
        
        if fridge_top_devices and fridge_top_devices[0]['alarm_count'] > 20:
            analysis_lines.append(f"→ 建议对高报警设备 '{fridge_top_devices[0]['name']}' 安排专项检查")
        
        analysis_lines.append("")
        analysis_lines.append("注：本分析由本地规则引擎生成（AI服务暂不可用），仅供参考。")
        
        return '\n'.join(analysis_lines)
    
    def export_alarm_data(self, alarm_data: Dict, analysis_result: Optional[Dict], date_range: Dict):
        import os
        import json
        from datetime import datetime
        
        export_data = {
            "export_time": datetime.now().isoformat(),
            "date_range": date_range,
            "alarm_data": alarm_data,
            "analysis_result": analysis_result
        }
        
        os.makedirs("outputs", exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"alarm_export_{timestamp}.json"
        file_path = os.path.join("outputs", filename)
        
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(export_data, f, ensure_ascii=False, indent=4)
        
        return {
            "filename": filename,
            "file_path": file_path,
            "total_alarms": alarm_data.get("total", 0),
            "export_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }

    def export_charts_excel(self, alarm_data: Dict):
        try:
            import os
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.chart import BarChart, PieChart, Reference
            from openpyxl.chart.label import DataLabelList
            
            devices = alarm_data.get("devices", [])
            
            os.makedirs("outputs", exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"alarm_charts_{timestamp}.xlsx"
            file_path = os.path.join("outputs", filename)
            
            wb = Workbook()
            wb.remove(wb.active)
            
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            ws0 = wb.create_sheet("全部报警数据", 0)
            data_headers = ["序号", "设备编号", "设备名称", "报警时间", "报警信息", "处理状态", "处理备注", "备注", "报警触发"]
            ws0.append(data_headers)
            for cell in ws0[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.border = thin_border
            
            seq = 0
            for device in devices:
                device_sn = device.get("device_sn", "")
                device_name = device.get("device_name", "")
                for alarm in device.get("alarms", []):
                    seq += 1
                    row_data = [
                        seq,
                        device_sn,
                        device_name,
                        alarm.get("alarmdate", ""),
                        alarm.get("message", ""),
                        alarm.get("handlestate", "未处理"),
                        alarm.get("handleremark", ""),
                        alarm.get("remark", ""),
                        alarm.get("alarmtrigger", "")
                    ]
                    ws0.append(row_data)
            
            for row in ws0.iter_rows(min_row=1, max_row=ws0.max_row, max_col=len(data_headers)):
                for cell in row:
                    cell.border = thin_border
            
            ws1 = wb.create_sheet("设备报警统计")
            ws1.append(["设备编号", "设备名称", "报警次数"])
            for cell in ws1[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center')
            
            for device in devices:
                ws1.append([
                    device.get("device_sn", ""),
                    device.get("device_name", ""),
                    len(device.get("alarms", []))
                ])
            
            if len(devices) > 0:
                chart1 = BarChart()
                chart1.type = "col"
                chart1.style = 10
                chart1.title = "设备报警统计"
                chart1.y_axis.title = "报警次数"
                chart1.x_axis.title = "设备编号"
                data1 = Reference(ws1, min_col=3, min_row=1, max_row=len(devices)+1)
                cats1 = Reference(ws1, min_col=1, min_row=2, max_row=len(devices)+1)
                chart1.add_data(data1, titles_from_data=True)
                chart1.set_categories(cats1)
                chart1.height = 12
                chart1.width = 22
                ws1.add_chart(chart1, "E2")
            
            ws2 = wb.create_sheet("报警类型分布")
            ws2.append(["报警类型", "数量", "占比"])
            for cell in ws2[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center')
            
            type_counts = {"温度报警": 0, "湿度报警": 0, "其他报警": 0}
            for device in devices:
                for alarm in device.get("alarms", []):
                    msg = alarm.get("message", "")
                    has_temp = "温度" in msg
                    has_hum = "湿度" in msg
                    if has_temp:
                        type_counts["温度报警"] += 1
                    if has_hum:
                        type_counts["湿度报警"] += 1
                    if not has_temp and not has_hum:
                        type_counts["其他报警"] += 1
            
            type_counts = {k: v for k, v in type_counts.items() if v > 0}
            
            total = sum(type_counts.values()) if type_counts else 1
            for type_name, count in type_counts.items():
                ws2.append([type_name, count, f"{(count/total*100):.1f}%"])
            
            if type_counts:
                chart2 = PieChart()
                labels2 = Reference(ws2, min_col=1, min_row=2, max_row=len(type_counts)+1)
                data2 = Reference(ws2, min_col=2, min_row=1, max_row=len(type_counts)+1)
                chart2.add_data(data2, titles_from_data=True)
                chart2.set_categories(labels2)
                chart2.title = "报警类型分布"
                chart2.dataLabels = DataLabelList()
                chart2.dataLabels.showPercent = True
                chart2.dataLabels.showLeaderLines = True
                chart2.height = 12
                chart2.width = 22
                ws2.add_chart(chart2, "E2")
            
            ws3 = wb.create_sheet("24小时分布")
            ws3.append(["小时", "温度报警", "湿度报警", "其他报警"])
            for cell in ws3[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center')
            
            hour_data = {h: {"temp": 0, "hum": 0, "other": 0} for h in range(24)}
            for device in devices:
                for alarm in device.get("alarms", []):
                    try:
                        dt = datetime.fromisoformat(alarm.get("alarmdate", "").replace("Z", "+00:00"))
                        h = dt.hour
                    except:
                        h = 0
                    msg = alarm.get("message", "")
                    has_temp = "温度" in msg
                    has_hum = "湿度" in msg
                    if has_temp:
                        hour_data[h]["temp"] += 1
                    if has_hum:
                        hour_data[h]["hum"] += 1
                    if not has_temp and not has_hum:
                        hour_data[h]["other"] += 1
            
            for h in range(24):
                ws3.append([
                    f"{h}:00",
                    hour_data[h]["temp"],
                    hour_data[h]["hum"],
                    hour_data[h]["other"]
                ])
            
            chart3 = BarChart()
            chart3.type = "col"
            chart3.style = 10
            chart3.title = "24小时报警分布"
            chart3.y_axis.title = "报警次数"
            chart3.x_axis.title = "小时"
            data3 = Reference(ws3, min_col=2, min_row=1, max_col=4, max_row=25)
            cats3 = Reference(ws3, min_col=1, min_row=2, max_row=25)
            chart3.add_data(data3, titles_from_data=True)
            chart3.set_categories(cats3)
            chart3.height = 12
            chart3.width = 22
            ws3.add_chart(chart3, "F2")
            
            ws4 = wb.create_sheet("星期分布")
            ws4.append(["星期", "温度报警", "湿度报警", "其他报警"])
            for cell in ws4[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center')
            
            weekdays = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"]
            week_data = {d: {"temp": 0, "hum": 0, "other": 0} for d in range(7)}
            for device in devices:
                for alarm in device.get("alarms", []):
                    try:
                        dt = datetime.fromisoformat(alarm.get("alarmdate", "").replace("Z", "+00:00"))
                        d = dt.weekday()
                    except:
                        d = 0
                    msg = alarm.get("message", "")
                    has_temp = "温度" in msg
                    has_hum = "湿度" in msg
                    if has_temp:
                        week_data[d]["temp"] += 1
                    if has_hum:
                        week_data[d]["hum"] += 1
                    if not has_temp and not has_hum:
                        week_data[d]["other"] += 1
            
            for d in range(7):
                ws4.append([
                    weekdays[d],
                    week_data[d]["temp"],
                    week_data[d]["hum"],
                    week_data[d]["other"]
                ])
            
            chart4 = BarChart()
            chart4.title = "星期报警分布"
            chart4.type = "col"
            chart4.style = 10
            chart4.y_axis.title = "报警次数"
            chart4.x_axis.title = "星期"
            data4 = Reference(ws4, min_col=2, min_row=1, max_col=4, max_row=8)
            cats4 = Reference(ws4, min_col=1, min_row=2, max_row=8)
            chart4.add_data(data4, titles_from_data=True)
            chart4.set_categories(cats4)
            chart4.height = 12
            chart4.width = 22
            ws4.add_chart(chart4, "F2")
            
            ws5 = wb.create_sheet("月度日期分布")
            ws5.append(["日期", "温度报警", "湿度报警", "其他报警"])
            for cell in ws5[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal='center', vertical='center')
            
            day_data = {d: {"temp": 0, "hum": 0, "other": 0} for d in range(1, 32)}
            for device in devices:
                for alarm in device.get("alarms", []):
                    try:
                        dt = datetime.fromisoformat(alarm.get("alarmdate", "").replace("Z", "+00:00"))
                        d = dt.day
                    except:
                        d = 1
                    msg = alarm.get("message", "")
                    has_temp = "温度" in msg
                    has_hum = "湿度" in msg
                    if has_temp:
                        day_data[d]["temp"] += 1
                    if has_hum:
                        day_data[d]["hum"] += 1
                    if not has_temp and not has_hum:
                        day_data[d]["other"] += 1
            
            for d in range(1, 32):
                ws5.append([
                    f"{d}日",
                    day_data[d]["temp"],
                    day_data[d]["hum"],
                    day_data[d]["other"]
                ])
            
            chart5 = BarChart()
            chart5.type = "col"
            chart5.style = 10
            chart5.title = "月度日期报警分布"
            chart5.y_axis.title = "报警次数"
            chart5.x_axis.title = "日期"
            data5 = Reference(ws5, min_col=2, min_row=1, max_col=4, max_row=32)
            cats5 = Reference(ws5, min_col=1, min_row=2, max_row=32)
            chart5.add_data(data5, titles_from_data=True)
            chart5.set_categories(cats5)
            chart5.height = 12
            chart5.width = 22
            ws5.add_chart(chart5, "F2")
            
            for ws in wb.worksheets:
                for col in ws.columns:
                    max_length = 0
                    column = col[0].column_letter
                    for cell in col:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = (max_length + 2) * 1.5
                    ws.column_dimensions[column].width = adjusted_width
            
            wb.save(file_path)
            return {
                "filename": filename,
                "file_path": file_path,
                "export_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
        except Exception as e:
            logger.error(f"导出Excel失败: {e}", exc_info=True)
            raise

    def export_ai_txt(self, analysis_data: Dict):
        try:
            import os
            os.makedirs("outputs", exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"ai_analysis_{timestamp}.txt"
            file_path = os.path.join("outputs", filename)
            
            content = []
            content.append("=" * 80)
            content.append("AI深度分析报告")
            content.append("=" * 80)
            content.append("")
            
            overview = analysis_data.get("overview", {})
            if overview:
                content.append("【一、报警概况总结】")
                content.append("-" * 40)
                content.append(f"总报警次数: {overview.get('total_alarms', 0)}")
                content.append(f"处理完成率: {overview.get('handled_rate', 0)}%")
                content.append(f"环比变化: {overview.get('trend_change', '--')}")
                content.append(f"统计周期: {overview.get('period', '--')}")
                content.append("")
            
            device_analysis = analysis_data.get("device_analysis", {})
            if device_analysis:
                content.append("【二、设备报警规律分析】")
                content.append("-" * 40)
                
                top_devices = device_analysis.get("top_devices", [])
                if top_devices:
                    content.append("1. 报警设备TOP排行:")
                    for i, dev in enumerate(top_devices, 1):
                        content.append(f"   {i}. {dev.get('name', '')} - {dev.get('alarm_count', 0)}次")
                    content.append("")
                
                critical_devices = device_analysis.get("critical_devices", [])
                if critical_devices:
                    content.append("2. 重点关注设备:")
                    for dev in critical_devices:
                        content.append(f"   - {dev}")
                    content.append("")
                
                peak_hours = device_analysis.get("peak_hours", [])
                if peak_hours:
                    content.append("3. 报警时段高峰:")
                    content.append(f"   {', '.join(peak_hours)}")
                    content.append("")
                
                peak_days = device_analysis.get("peak_days", [])
                if peak_days:
                    content.append("4. 报警周规律:")
                    content.append(f"   {', '.join(peak_days)}")
                    content.append("")
                
                peak_month_days = device_analysis.get("peak_month_days", [])
                if peak_month_days:
                    content.append("5. 报警月规律:")
                    content.append(f"   {', '.join(peak_month_days)}")
                    content.append("")
            
            type_analysis = analysis_data.get("type_analysis", {})
            if type_analysis:
                content.append("【三、报警类型深度分析】")
                content.append("-" * 40)
                
                type_dist = type_analysis.get("distribution", [])
                if type_dist:
                    content.append("1. 报警类型分布:")
                    for item in type_dist:
                        content.append(f"   - {item.get('type', '')}: {item.get('count', 0)}次, 占比 {item.get('percentage', 0)}%")
                    content.append("")
            
            analysis_text = analysis_data.get("analysis_text", "")
            if analysis_text:
                content.append("【四、分析报告全文】")
                content.append("-" * 40)
                content.append(analysis_text)
                content.append("")
            
            content.append("=" * 80)
            content.append(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            content.append("=" * 80)
            
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("\n".join(content))
            
            return {
                "filename": filename,
                "file_path": file_path,
                "export_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
        except Exception as e:
            logger.error(f"导出TXT失败: {e}", exc_info=True)
            raise

    def export_ai_word(self, analysis_data: Dict):
        try:
            import os
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_ALIGN_VERTICAL
            from docx.oxml.ns import qn
            
            def set_chinese_font(run):
                """统一设置中文字体为宋体"""
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            os.makedirs("outputs", exist_ok=True)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"ai_analysis_{timestamp}.docx"
            file_path = os.path.join("outputs", filename)
            
            doc = Document()
            
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            font.size = Pt(11)
            
            title = doc.add_heading("", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("AI深度分析报告")
            set_chinese_font(run)
            run.font.size = Pt(20)
            run.font.bold = True
            
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title_para.add_run(f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            run.font.size = Pt(10)
            set_chinese_font(run)
            
            doc.add_paragraph()
            
            doc.add_heading("一、报警概况总结", level=1)
            overview = analysis_data.get("overview", {})
            if overview:
                table = doc.add_table(rows=1, cols=4)
                table.style = "Table Grid"
                table.autofit = True
                hdr_cells = table.rows[0].cells
                for i, (cell, header_text) in enumerate(zip(hdr_cells, ["总报警次数", "处理完成率", "环比变化", "统计周期"])):
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(header_text)
                    run.font.bold = True
                    set_chinese_font(run)
                row_cells = table.add_row().cells
                for cell in row_cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph = cell.paragraphs[0]
                    paragraph.clear()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(cell.text if cell.text else "")
                    set_chinese_font(run)
                para0 = row_cells[0].paragraphs[0]
                para0.clear()
                para0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run0 = para0.add_run(str(overview.get("total_alarms", 0)))
                set_chinese_font(run0)
                para1 = row_cells[1].paragraphs[0]
                para1.clear()
                para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run1 = para1.add_run(f"{overview.get('handled_rate', 0)}%")
                set_chinese_font(run1)
                para2 = row_cells[2].paragraphs[0]
                para2.clear()
                para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run2 = para2.add_run(str(overview.get("trend_change", "--")))
                set_chinese_font(run2)
                para3 = row_cells[3].paragraphs[0]
                para3.clear()
                para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run3 = para3.add_run(str(overview.get("period", "--")))
                set_chinese_font(run3)
            
            doc.add_heading("二、设备报警规律分析", level=1)
            device_analysis = analysis_data.get("device_analysis", {})
            if device_analysis:
                top_devices = device_analysis.get("top_devices", [])
                if top_devices:
                    doc.add_heading("1. 报警设备TOP排行", level=2)
                    table = doc.add_table(rows=1, cols=3)
                    table.style = "Table Grid"
                    table.autofit = True
                    hdr_cells = table.rows[0].cells
                    for i, (cell, header_text) in enumerate(zip(hdr_cells, ["排名", "设备名称", "报警次数"])):
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run(header_text)
                        run.font.bold = True
                        set_chinese_font(run)
                    for i, dev in enumerate(top_devices, 1):
                        row_cells = table.add_row().cells
                        for cell in row_cells:
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph = cell.paragraphs[0]
                            paragraph.clear()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run("")
                            set_chinese_font(run)
                        para0 = row_cells[0].paragraphs[0]
                        para0.clear()
                        para0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run0 = para0.add_run(str(i))
                        set_chinese_font(run0)
                        para1 = row_cells[1].paragraphs[0]
                        para1.clear()
                        para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run1 = para1.add_run(str(dev.get("name", "")))
                        set_chinese_font(run1)
                        para2 = row_cells[2].paragraphs[0]
                        para2.clear()
                        para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run2 = para2.add_run(str(dev.get("alarm_count", 0)))
                        set_chinese_font(run2)
                
                critical_devices = device_analysis.get("critical_devices", [])
                if critical_devices:
                    doc.add_heading("2. 重点关注设备", level=2)
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.5
                    for dev in critical_devices:
                        run = p.add_run(f"• {dev}\n")
                        set_chinese_font(run)
                
                peak_hours = device_analysis.get("peak_hours", [])
                if peak_hours:
                    doc.add_heading("3. 报警时段高峰", level=2)
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.5
                    run = p.add_run(f"报警高发时段：{', '.join(peak_hours)}")
                    set_chinese_font(run)
                
                peak_days = device_analysis.get("peak_days", [])
                if peak_days:
                    doc.add_heading("4. 报警周规律", level=2)
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.5
                    run = p.add_run(f"报警高发星期：{', '.join(peak_days)}")
                    set_chinese_font(run)
                
                peak_month_days = device_analysis.get("peak_month_days", [])
                if peak_month_days:
                    doc.add_heading("5. 报警月规律", level=2)
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.5
                    run = p.add_run(f"报警高发日期：{', '.join(peak_month_days)}")
                    set_chinese_font(run)
            
            doc.add_heading("三、报警类型深度分析", level=1)
            type_analysis = analysis_data.get("type_analysis", {})
            if type_analysis:
                type_dist = type_analysis.get("distribution", [])
                if type_dist:
                    doc.add_heading("1. 报警类型分布", level=2)
                    table = doc.add_table(rows=1, cols=3)
                    table.style = "Table Grid"
                    table.autofit = True
                    hdr_cells = table.rows[0].cells
                    for i, (cell, header_text) in enumerate(zip(hdr_cells, ["报警类型", "数量", "占比"])):
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run(header_text)
                        run.font.bold = True
                        set_chinese_font(run)
                    for item in type_dist:
                        row_cells = table.add_row().cells
                        for cell in row_cells:
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph = cell.paragraphs[0]
                            paragraph.clear()
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run("")
                            set_chinese_font(run)
                        para0 = row_cells[0].paragraphs[0]
                        para0.clear()
                        para0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run0 = para0.add_run(str(item.get('type', '')))
                        set_chinese_font(run0)
                        para1 = row_cells[1].paragraphs[0]
                        para1.clear()
                        para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run1 = para1.add_run(str(item.get('count', 0)))
                        set_chinese_font(run1)
                        para2 = row_cells[2].paragraphs[0]
                        para2.clear()
                        para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run2 = para2.add_run(f"{item.get('percentage', 0)}%")
                        set_chinese_font(run2)
                
                severity_stats = type_analysis.get("severity_statistics", {})
                if severity_stats:
                    doc.add_heading("2. 严重程度分布", level=2)
                    table = doc.add_table(rows=1, cols=3)
                    table.style = "Table Grid"
                    table.autofit = True
                    hdr_cells = table.rows[0].cells
                    for i, (cell, header_text) in enumerate(zip(hdr_cells, ["高优先级", "中优先级", "低优先级"])):
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run(header_text)
                        run.font.bold = True
                        set_chinese_font(run)
                    row_cells = table.add_row().cells
                    for cell in row_cells:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        paragraph = cell.paragraphs[0]
                        paragraph.clear()
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = paragraph.add_run("")
                        set_chinese_font(run)
                    para0 = row_cells[0].paragraphs[0]
                    para0.clear()
                    para0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run0 = para0.add_run(str(severity_stats.get("高", 0)))
                    set_chinese_font(run0)
                    para1 = row_cells[1].paragraphs[0]
                    para1.clear()
                    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run1 = para1.add_run(str(severity_stats.get("中", 0)))
                    set_chinese_font(run1)
                    para2 = row_cells[2].paragraphs[0]
                    para2.clear()
                    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run2 = para2.add_run(str(severity_stats.get("低", 0)))
                    set_chinese_font(run2)
                
                by_category = type_analysis.get("by_category", {})
                if by_category:
                    doc.add_heading("3. 设备分类排行", level=2)
                    
                    fridge_data = by_category.get("fridge", {})
                    if fridge_data:
                        doc.add_heading("(1) 冰箱温度设备", level=3)
                        table = doc.add_table(rows=1, cols=3)
                        table.style = "Table Grid"
                        table.autofit = True
                        hdr_cells = table.rows[0].cells
                        for i, (cell, header_text) in enumerate(zip(hdr_cells, ["排名", "设备名称", "报警次数"])):
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph = cell.paragraphs[0]
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run(header_text)
                            run.font.bold = True
                            set_chinese_font(run)
                        fridge_top = fridge_data.get("top_devices", [])
                        for i, dev in enumerate(fridge_top[:5], 1):
                            row_cells = table.add_row().cells
                            for cell in row_cells:
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                paragraph = cell.paragraphs[0]
                                paragraph.clear()
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = paragraph.add_run("")
                                set_chinese_font(run)
                            para0 = row_cells[0].paragraphs[0]
                            para0.clear()
                            para0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run0 = para0.add_run(str(i))
                            set_chinese_font(run0)
                            para1 = row_cells[1].paragraphs[0]
                            para1.clear()
                            para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run1 = para1.add_run(str(dev.get("name", "")))
                            set_chinese_font(run1)
                            para2 = row_cells[2].paragraphs[0]
                            para2.clear()
                            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run2 = para2.add_run(str(dev.get("alarm_count", 0)))
                            set_chinese_font(run2)
                        p = doc.add_paragraph()
                        p.paragraph_format.line_spacing = 1.5
                        run = p.add_run(f"设备总数：{len(fridge_data.get('devices', []))}台，报警总数：{fridge_data.get('total_alarms', 0)}次")
                        set_chinese_font(run)
                    
                    env_data = by_category.get("environment", {})
                    if env_data:
                        doc.add_heading("(2) 环境温湿度设备", level=3)
                        table = doc.add_table(rows=1, cols=3)
                        table.style = "Table Grid"
                        table.autofit = True
                        hdr_cells = table.rows[0].cells
                        for i, (cell, header_text) in enumerate(zip(hdr_cells, ["排名", "设备名称", "报警次数"])):
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            paragraph = cell.paragraphs[0]
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = paragraph.add_run(header_text)
                            run.font.bold = True
                            set_chinese_font(run)
                        env_top = env_data.get("top_devices", [])
                        for i, dev in enumerate(env_top[:5], 1):
                            row_cells = table.add_row().cells
                            for cell in row_cells:
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                paragraph = cell.paragraphs[0]
                                paragraph.clear()
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = paragraph.add_run("")
                                set_chinese_font(run)
                            para0 = row_cells[0].paragraphs[0]
                            para0.clear()
                            para0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run0 = para0.add_run(str(i))
                            set_chinese_font(run0)
                            para1 = row_cells[1].paragraphs[0]
                            para1.clear()
                            para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run1 = para1.add_run(str(dev.get("name", "")))
                            set_chinese_font(run1)
                            para2 = row_cells[2].paragraphs[0]
                            para2.clear()
                            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run2 = para2.add_run(str(dev.get("alarm_count", 0)))
                            set_chinese_font(run2)
                        p = doc.add_paragraph()
                        p.paragraph_format.line_spacing = 1.5
                        run = p.add_run(f"设备总数：{len(env_data.get('devices', []))}台，报警总数：{env_data.get('total_alarms', 0)}次")
                        set_chinese_font(run)
            
            doc.add_heading("四、AI分析建议", level=1)
            suggestions = analysis_data.get("suggestions", [])
            if suggestions:
                for i, suggestion in enumerate(suggestions, 1):
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.5
                    run = p.add_run(f"{i}. 【{suggestion.get('priority', '')}优先级】{suggestion.get('type', '')}\n")
                    run.bold = True
                    set_chinese_font(run)
                    run = p.add_run(f"   针对: {suggestion.get('target', '')}\n")
                    set_chinese_font(run)
                    run = p.add_run(f"   建议: {suggestion.get('description', '')}\n")
                    set_chinese_font(run)
            
            doc.add_heading("五、处理情况评估", level=1)
            handling_eval = analysis_data.get("handling_evaluation", {})
            if handling_eval:
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                table.autofit = True
                hdr_cells = table.rows[0].cells
                for i, (cell, header_text) in enumerate(zip(hdr_cells, ["未处理报警", "平均响应时间"])):
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run(header_text)
                    run.font.bold = True
                    set_chinese_font(run)
                row_cells = table.add_row().cells
                for cell in row_cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    paragraph = cell.paragraphs[0]
                    paragraph.clear()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run("")
                    set_chinese_font(run)
                para0 = row_cells[0].paragraphs[0]
                para0.clear()
                para0.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run0 = para0.add_run(f"{handling_eval.get('unhandled_count', 0)}条")
                set_chinese_font(run0)
                para1 = row_cells[1].paragraphs[0]
                para1.clear()
                para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run1 = para1.add_run(str(handling_eval.get("average_response_time", "--")))
                set_chinese_font(run1)
            
            doc.add_heading("六、分析报告全文", level=1)
            analysis_text = analysis_data.get("analysis_text", "")
            if analysis_text:
                for line in analysis_text.split("\n"):
                    p = doc.add_paragraph()
                    p.paragraph_format.line_spacing = 1.5
                    run = p.add_run(line)
                    set_chinese_font(run)
            
            # 设置真正的页脚
            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 清除原有内容并重新设置
            for element in footer_para._element:
                footer_para._element.remove(element)
            run = footer_para.add_run(f"TraceTempAI AI深度分析报告 - 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            run.font.size = Pt(10)
            set_chinese_font(run)
            
            doc.save(file_path)
            return {
                "filename": filename,
                "file_path": file_path,
                "export_time": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
        except Exception as e:
            logger.error(f"导出Word失败: {e}", exc_info=True)
            raise

class ReportService:
    def __init__(self):
        self.dept_manager = DepartmentManager()
        self.task_manager = TaskManager()
    
    def generate_report(self, report_type: str, start_date: str, end_date: str, dept_id: Optional[str] = None):
        if report_type not in ["monthly_review", "correction"]:
            raise ValueError(f"不支持的报告类型: {report_type}")
        
        if dept_id:
            dept = self.dept_manager.get_department(dept_id)
            if not dept:
                raise ValueError(f"学科 {dept_id} 不存在")
            api_key = dept.api_key
            dept_name = dept.name
        else:
            departments = self.dept_manager.get_all_departments()
            if not departments:
                raise ValueError("未配置任何学科")
            dept = departments[0]
            api_key = dept.api_key
            dept_name = dept.name
        
        # 使用可变容器保存 task_id，report_task 在后台线程中通过它更新进度
        task_holder = {"task_id": None}
        
        def report_task():
            import os
            task_id_local = task_holder.get("task_id")
            start_dt = datetime.strptime(start_date, "%Y-%m-%d")
            end_dt = datetime.strptime(end_date, "%Y-%m-%d").replace(hour=23, minute=59, second=59)
            
            # 阶段1：数据采集（0-70%）
            def collect_progress(current, total):
                if total > 0 and task_id_local:
                    pct = int(current / total * 70)
                    self.task_manager.update_task_progress(task_id_local, pct)
            
            collector = DataCollector(api_key=api_key)
            alarm_data = collector.get_all_alarms(start_dt, end_dt, progress_callback=collect_progress)
            
            # 阶段2：报告生成（70-90%）
            if task_id_local:
                self.task_manager.update_task_progress(task_id_local, 72)
            generator = ReportGenerator()
            if report_type == "monthly_review":
                result = generator.generate_monthly_report(alarm_data, dept_name)
                results = [result]
            else:
                results = generator.generate_correction_report(alarm_data, dept_name)
            if task_id_local:
                self.task_manager.update_task_progress(task_id_local, 90)
            
            # 阶段3：保存文件（90-100%）
            os.makedirs("outputs/reports", exist_ok=True)
            saved_results = []
            total_files = len(results) if results else 1
            for idx, r in enumerate(results):
                file_path = os.path.join("outputs/reports", r["filename"])
                with open(file_path, "wb") as f:
                    f.write(r["content"])
                saved_results.append({
                    "filename": r["filename"],
                    "file_path": file_path,
                    "content": r["content"]
                })
                if task_id_local:
                    pct = 90 + int((idx + 1) / total_files * 9)
                    self.task_manager.update_task_progress(task_id_local, min(99, pct))
            
            return saved_results
        
        task_id = self.task_manager.create_task(
            name=f"生成{report_type}",
            func=report_task
        )
        task_holder["task_id"] = task_id
        self.task_manager.tasks[task_id].report_type = report_type
        self.task_manager.tasks[task_id].department = dept_name
        
        return task_id
    
    def get_tasks(self):
        from datetime import datetime
        tasks = self.task_manager.get_all_tasks()
        result = []
        for task_id, task in tasks.items():
            # 从 result 中获取所有 file_ids
            file_ids = []
            if task.result and isinstance(task.result, list) and len(task.result) > 0:
                file_ids = [r.get("filename") for r in task.result if r.get("filename")]
            
            result.append({
                "task_id": task.task_id,
                "name": task.name,
                "status": task.status.value,
                "progress": task.progress,
                "created_at": datetime.fromtimestamp(task.created_at).strftime("%Y-%m-%d %H:%M:%S") if task.created_at else None,
                "started_at": datetime.fromtimestamp(task.started_at).strftime("%Y-%m-%d %H:%M:%S") if task.started_at else None,
                "completed_at": datetime.fromtimestamp(task.completed_at).strftime("%Y-%m-%d %H:%M:%S") if task.completed_at else None,
                "report_type": task.report_type,
                "department": task.department,
                "error": task.error,
                "file_ids": file_ids
            })
        result.sort(key=lambda x: x["created_at"] or "", reverse=True)
        return result
    
    def get_task_status(self, task_id: str):
        from datetime import datetime
        task = self.task_manager.get_task(task_id)
        if not task:
            raise ValueError(f"任务 {task_id} 不存在")
        
        # 从 result 中获取所有 file_ids
        file_ids = []
        if task.result and isinstance(task.result, list) and len(task.result) > 0:
            file_ids = [r.get("filename") for r in task.result if r.get("filename")]
        
        return {
            "task_id": task.task_id,
            "name": task.name,
            "status": task.status.value,
            "progress": task.progress,
            "created_at": datetime.fromtimestamp(task.created_at).strftime("%Y-%m-%d %H:%M:%S") if task.created_at else None,
            "started_at": datetime.fromtimestamp(task.started_at).strftime("%Y-%m-%d %H:%M:%S") if task.started_at else None,
            "completed_at": datetime.fromtimestamp(task.completed_at).strftime("%Y-%m-%d %H:%M:%S") if task.completed_at else None,
            "report_type": task.report_type,
            "department": task.department,
            "error": task.error,
            "result": task.result,
            "file_ids": file_ids
        }
    
    def download_reports_batch(self, file_ids: list):
        """
        批量下载报告文件，打包为 ZIP

        包含路径遍历防护：对每个 file_id 校验其解析后的真实路径
        是否位于 outputs/reports 目录内，防止目录穿越攻击。

        Args:
            file_ids: 文件名列表（仅允许纯文件名，禁止路径分隔符）

        Returns:
            dict: 包含 zip_filename、file_path、count 的结果字典
        """
        import os
        import zipfile
        from datetime import datetime
        import time
        
        os.makedirs("outputs", exist_ok=True)
        base_dir = os.path.realpath("outputs/reports")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        zip_filename = f"reports_batch_{timestamp}.zip"
        zip_path = os.path.join("outputs", zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_id in file_ids:
                safe_filename = os.path.basename(file_id)
                file_path = os.path.realpath(os.path.join("outputs/reports", safe_filename))
                if file_path.startswith(base_dir + os.sep) and os.path.isfile(file_path):
                    # 手动创建ZipInfo，确保文件名使用UTF-8编码
                    zip_info = zipfile.ZipInfo(safe_filename)
                    zip_info.date_time = time.localtime(time.time())[:6]
                    zip_info.compress_type = zipfile.ZIP_DEFLATED
                    # 设置UTF-8标志（bit 11），确保文件名正确编码
                    zip_info.flag_bits |= 0x800
                    
                    with open(file_path, 'rb') as f:
                        zipf.writestr(zip_info, f.read())
        
        return {
            "filename": zip_filename,
            "file_path": zip_path,
            "count": len(file_ids)
        }

class DepartmentService:
    def __init__(self):
        self.dept_manager = DepartmentManager()
    
    def get_all_departments(self):
        departments = self.dept_manager.get_all_departments()
        return [{"id": dept.id, "name": dept.name} for dept in departments]